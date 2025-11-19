"""
Document Processing Service
Handles extraction of text from various file formats:
- PDFs
- Images (OCR)
- Word documents
"""

import io
import logging
import os
from pathlib import Path
from typing import Optional

import cv2
import numpy as np
from PIL import Image

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process and extract text from various document formats"""

    def __init__(self):
        self.ocr_available = False
        self.pdf_available = False
        self.docx_available = False

        # Try to import OCR libraries
        try:
            import easyocr

            self.reader = easyocr.Reader(["en"], gpu=False)
            self.ocr_available = True
            logger.info("✅ EasyOCR initialized")
        except Exception as e:
            logger.warning(f"⚠️ EasyOCR not available: {e}")
            try:
                import pytesseract

                # Set Tesseract path for Windows
                tesseract_paths = [
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                    r"C:\Users\{}\AppData\Local\Tesseract-OCR\tesseract.exe".format(
                        os.getenv("USERNAME")
                    ),
                ]

                for path in tesseract_paths:
                    if Path(path).exists():
                        pytesseract.pytesseract.tesseract_cmd = path
                        logger.info(f"✅ Tesseract found at: {path}")
                        break

                self.ocr_available = True
                logger.info("✅ Tesseract OCR initialized")
            except Exception as e2:
                logger.warning(f"⚠️ Tesseract not available: {e2}")

        # Try to import PDF libraries
        try:
            import pdfplumber
            import PyPDF2

            self.pdf_available = True
            logger.info("✅ PDF processing available")
        except Exception as e:
            logger.warning(f"⚠️ PDF processing not available: {e}")

        # Try to import DOCX library
        try:
            import docx

            self.docx_available = True
            logger.info("✅ DOCX processing available")
        except Exception as e:
            logger.warning(f"⚠️ DOCX processing not available: {e}")

    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from file based on extension

        Args:
            file_content: Raw file bytes
            filename: Name of the file with extension

        Returns:
            Extracted text content
        """
        ext = Path(filename).suffix.lower()

        try:
            if ext == ".pdf":
                return await self._extract_from_pdf(file_content)
            elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
                return await self._extract_from_image(file_content)
            elif ext in [".doc", ".docx"]:
                return await self._extract_from_docx(file_content)
            elif ext == ".txt":
                return file_content.decode("utf-8", errors="ignore")
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            logger.error(f"❌ Error extracting text from {filename}: {e}")
            raise

    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        if not self.pdf_available:
            raise RuntimeError("PDF processing libraries not available")

        from io import BytesIO

        import pdfplumber

        text_parts = []

        try:
            # Try pdfplumber first (better for structured PDFs)
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"\n--- Page {page_num} ---\n")
                        text_parts.append(page_text)

            if text_parts:
                return "".join(text_parts).strip()
        except Exception as e:
            logger.warning(f"⚠️ pdfplumber failed: {e}, trying PyPDF2")

        # Fallback to PyPDF2
        try:
            from io import BytesIO

            import PyPDF2

            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))

            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"\n--- Page {page_num} ---\n")
                    text_parts.append(page_text)

            if text_parts:
                return "".join(text_parts).strip()
        except Exception as e:
            logger.error(f"❌ PyPDF2 also failed: {e}")

        # If text extraction failed, PDF might be image-based
        # Try OCR on PDF pages
        if self.ocr_available:
            try:
                from pdf2image import convert_from_bytes

                logger.info("📄 PDF appears to be image-based, using OCR...")
                images = convert_from_bytes(file_content)

                for page_num, img in enumerate(images, 1):
                    # Convert PIL Image to numpy array for OCR
                    img_array = np.array(img)
                    page_text = await self._ocr_image(img_array)
                    if page_text:
                        text_parts.append(f"\n--- Page {page_num} (OCR) ---\n")
                        text_parts.append(page_text)

                if text_parts:
                    return "".join(text_parts).strip()
            except Exception as ocr_error:
                logger.error(f"❌ PDF OCR failed: {ocr_error}")

        raise RuntimeError(
            "Could not extract text from PDF. File may be corrupted or password-protected."
        )

    async def _extract_from_image(self, file_content: bytes) -> str:
        """Extract text from image using OCR"""
        if not self.ocr_available:
            raise RuntimeError("OCR libraries not available")

        # Convert bytes to numpy array
        nparr = np.frombuffer(file_content, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

        if img is None:
            raise ValueError("Invalid image file")

        return await self._ocr_image(img)

    async def _ocr_image(self, img: np.ndarray) -> str:
        """Perform OCR on image array with AI fallback"""
        # Preprocess image for better OCR
        img = self._preprocess_image(img)

        try:
            # Try EasyOCR first
            if hasattr(self, "reader"):
                results = self.reader.readtext(img)
                text = " ".join([result[1] for result in results])
                return text.strip()
        except Exception as e:
            logger.warning(f"⚠️ EasyOCR failed: {e}, trying Tesseract")

        # Fallback to Tesseract
        try:
            import pytesseract
            from PIL import Image

            # Convert to PIL Image
            img_pil = Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
            text = pytesseract.image_to_string(img_pil)
            return text.strip()
        except Exception as e:
            logger.warning(f"⚠️ Tesseract OCR failed: {e}, trying AI vision fallback")

        # Final fallback: Manual extraction or error
        logger.error(
            "❌ All OCR methods failed. Tesseract is not installed. "
            "Please install Tesseract OCR or upload a PDF file instead."
        )
        raise RuntimeError(
            "Image OCR is not available on this server. "
            "Please upload a PDF file instead, or ensure the image contains clear, readable text. "
            "For best results, use PDF documents for medical reports."
        )

    def _preprocess_image(self, img: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR results"""
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Apply denoising
        denoised = cv2.fastNlMeansDenoising(gray)

        # Apply adaptive thresholding
        thresh = cv2.adaptiveThreshold(
            denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )

        return thresh

    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        if not self.docx_available:
            raise RuntimeError("DOCX processing library not available")

        from io import BytesIO

        import docx

        doc = docx.Document(BytesIO(file_content))

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts).strip()


# Global instance
document_processor = DocumentProcessor()
